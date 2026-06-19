import os
import io
import re
import tempfile
import subprocess
from typing import Optional

from fastapi import FastAPI, Header, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from docx import Document

API_KEY = os.environ.get("DOC_API_KEY", "")

app = FastAPI(title="AI Document Updater Service", version="1.0.0")


def _check_key(provided: Optional[str]):
    if not API_KEY:
        # If no key configured, refuse rather than run open
        raise HTTPException(status_code=500, detail="Service misconfigured: DOC_API_KEY not set")
    if provided != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid or missing X-API-Key")


# ----------------------------------------------------------------------
# Health
# ----------------------------------------------------------------------
@app.get("/health")
def health():
    # check libreoffice availability for PDF conversion
    lo = False
    for binname in ("libreoffice", "soffice"):
        try:
            subprocess.run([binname, "--version"], capture_output=True, timeout=10)
            lo = True
            break
        except Exception:
            continue
    return {"status": "ok", "libreoffice": lo}


# ----------------------------------------------------------------------
# Analyze: extract structured editable content from a DOCX
# ----------------------------------------------------------------------
DATE_RE = re.compile(r"\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b", re.I)
MONTH_RE = re.compile(r"\b(0?[1-9]|1[0-2])[/-]\d{4}\b")
YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")
CONTROL_RE = re.compile(r"\b(?:No\.?|Control(?:\s*No)?\.?|Ref(?:erence)?\.?)\s*[:#]?\s*[A-Z0-9\-/]+\b", re.I)


@app.post("/analyze")
async def analyze(file: UploadFile = File(...), x_api_key: Optional[str] = Header(None)):
    _check_key(x_api_key)
    data = await file.read()
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Cannot read DOCX: {e}")

    paragraphs = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            paragraphs.append({
                "index": idx,
                "style": p.style.name if p.style else None,
                "text": text,
            })

    tables = []
    for tidx, t in enumerate(doc.tables):
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        tables.append({"index": tidx, "rows": rows})

    full_text = "\n".join(p["text"] for p in paragraphs)
    dates = sorted(set(DATE_RE.findall(full_text)))
    months = sorted(set(m.group(0) for m in MONTH_RE.finditer(full_text)))
    years = sorted(set(YEAR_RE.findall(full_text)))
    controls = sorted(set(c.group(0) for c in CONTROL_RE.finditer(full_text)))

    # crude image count via relationships
    image_count = 0
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
    except Exception:
        pass

    return JSONResponse({
        "filename": file.filename,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "image_count": image_count,
        "detected": {
            "dates": dates,
            "month_codes": months,
            "years": years,
            "control_numbers": controls,
        },
        "paragraphs": paragraphs,
        "tables": tables,
    })


# ----------------------------------------------------------------------
# Update: apply text replacements while preserving run formatting
# ----------------------------------------------------------------------
def _replace_in_paragraph(paragraph, replacements):
    """Replace text across runs while keeping each run's formatting.
    Strategy: if the search term lives in a single run, edit that run.
    Otherwise rebuild the paragraph text into the first run (best-effort)."""
    for search, repl in replacements.items():
        if not search:
            continue
        # fast path: run-local replacement
        for run in paragraph.runs:
            if search in run.text:
                run.text = run.text.replace(search, repl)
        # cross-run path
        joined = "".join(r.text for r in paragraph.runs)
        if search in joined and paragraph.runs:
            new_joined = joined.replace(search, repl)
            # put everything in first run, clear the rest (keeps first run style)
            paragraph.runs[0].text = new_joined
            for r in paragraph.runs[1:]:
                r.text = ""


@app.post("/update")
async def update(
    file: UploadFile = File(...),
    replacements_json: str = Form(...),
    x_api_key: Optional[str] = Header(None),
):
    """replacements_json: a JSON object {"old text": "new text", ...}"""
    _check_key(x_api_key)
    import json
    try:
        replacements = json.loads(replacements_json)
        assert isinstance(replacements, dict)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"replacements_json must be a JSON object: {e}")

    data = await file.read()
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Cannot read DOCX: {e}")

    # body paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)
    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replacements)
    # headers/footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                _replace_in_paragraph(p, replacements)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="updated.docx"'},
    )


# ----------------------------------------------------------------------
# To-PDF: convert DOCX to PDF via LibreOffice headless
# ----------------------------------------------------------------------
@app.post("/to-pdf")
async def to_pdf(file: UploadFile = File(...), x_api_key: Optional[str] = Header(None)):
    _check_key(x_api_key)
    data = await file.read()
    with tempfile.TemporaryDirectory() as tmp:
        in_path = os.path.join(tmp, "in.docx")
        with open(in_path, "wb") as fh:
            fh.write(data)

        binname = None
        for b in ("libreoffice", "soffice"):
            try:
                subprocess.run([b, "--version"], capture_output=True, timeout=10)
                binname = b
                break
            except Exception:
                continue
        if not binname:
            raise HTTPException(status_code=500, detail="LibreOffice not available for PDF conversion")

        proc = subprocess.run(
            [binname, "--headless", "--convert-to", "pdf", "--outdir", tmp, in_path],
            capture_output=True, timeout=120,
        )
        pdf_path = os.path.join(tmp, "in.pdf")
        if not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail=f"PDF conversion failed: {proc.stderr.decode()[:300]}")
        with open(pdf_path, "rb") as fh:
            pdf_bytes = fh.read()

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="output.pdf"'},
    )
